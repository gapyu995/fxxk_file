import asyncio
import json
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

import app.main as main_module
import app.services.storage as storage_module
import app.services.translator as translator_module
from app.main import _bounded_item_batches, _docx_preview_page, _pdf_preview_page, _progressive_batches
from app.services.extractor import extract_paragraphs
from app.services.glossary import Term, relevant_terms
from app.services.segmenter import detect_language, split_into_segments
from app.services.translator import TranslationError, build_endpoint, translate_batch


def test_detect_language_both_directions():
    assert detect_language("这是一个专业翻译项目。") == "zh"
    assert detect_language("This is a professional translation project.") == "en"


def test_split_long_paragraph():
    result = split_into_segments(["第一句。第二句。第三句。"], max_chars=7)
    assert len(result) >= 2
    assert "".join(result).replace(" ", "") == "第一句。第二句。第三句。"


def test_relevant_terms_only_returns_matches():
    terms = [Term("人工智能", "artificial intelligence"), Term("合同", "contract")]
    matched = relevant_terms(terms, ["本文件讨论人工智能。"])
    assert [term.source for term in matched] == ["人工智能"]


def test_extract_docx(tmp_path: Path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("Second paragraph")
    document.save(path)
    assert extract_paragraphs(path) == ["第一段", "Second paragraph"]


def test_ccswitch_gateway_endpoint_variants():
    assert build_endpoint("https://llmapi.neuehct.auto", "openai") == (
        "https://llmapi.neuehct.auto/v1/chat/completions"
    )
    assert build_endpoint("https://gateway.example/v1", "openai") == (
        "https://gateway.example/v1/chat/completions"
    )
    assert build_endpoint("https://gateway.example/v1/chat/completions", "openai") == (
        "https://gateway.example/v1/chat/completions"
    )
    assert build_endpoint("https://gateway.example", "anthropic") == (
        "https://gateway.example/v1/messages"
    )


def test_progressive_batches_returns_first_segment_immediately():
    ids = [f"s{index}" for index in range(1, 6)]
    by_id = {item_id: {"source": "abc"} for item_id in ids}
    assert _progressive_batches(ids, by_id, max_count=3, max_chars=100) == [
        ["s1"],
        ["s2", "s3", "s4"],
        ["s5"],
    ]


def test_bounded_batches_respect_count_and_character_limits():
    items = [{"id": f"s{index}", "text": "abc"} for index in range(1, 6)]
    batches = _bounded_item_batches(items, max_count=3, max_chars=7)
    assert [[item["id"] for item in batch] for batch in batches] == [
        ["s1", "s2"],
        ["s3", "s4"],
        ["s5"],
    ]
    assert all(len(batch) <= 3 for batch in batches)
    assert all(sum(len(item["text"]) for item in batch) <= 7 for batch in batches)


def test_long_segment_is_split_across_requests_and_merged(monkeypatch):
    settings = SimpleNamespace(request_char_limit=6, batch_size=2)
    calls = []

    async def fake_translate_batch(
        items,
        source_lang,
        target_lang,
        terms,
        style_guide,
        received_settings,
        on_retry=None,
    ):
        calls.append(items)
        assert received_settings is settings
        return {item["id"]: "[" + item["text"] + "]" for item in items}

    monkeypatch.setattr(main_module, "get_settings", lambda: settings)
    monkeypatch.setattr(main_module, "load_terms", lambda *args: [])
    monkeypatch.setattr(main_module, "load_style_guide", lambda *args: "")
    monkeypatch.setattr(main_module, "translate_batch", fake_translate_batch)

    result = asyncio.run(
        main_module._translate_items(
            {"source_language": "en", "target_language": "zh"},
            [{"id": "s1", "text": "abcdefghijklmn"}],
        )
    )

    assert result == {"s1": "[abcdef][ghijkl][mn]"}
    assert len(calls) == 3
    assert all(len(batch) <= settings.batch_size for batch in calls)
    assert all(
        sum(len(item["text"]) for item in batch) <= settings.request_char_limit
        for batch in calls
    )


def test_preview_pages_include_bidirectional_sync_protocol():
    docx_html = _docx_preview_page("0123456789ab")
    pdf_html = _pdf_preview_page("0123456789ab")
    for html in (docx_html, pdf_html):
        assert "chanslator-sync-scroll" in html
        assert "chanslator-preview-scroll" in html
        assert "chanslator-preview-ready" in html
    assert "IntersectionObserver" in pdf_html
    assert "/vendor/pdf.worker.min.js" in pdf_html


def test_shutdown_clears_all_document_json_records(tmp_path: Path, monkeypatch):
    active = tmp_path / "0123456789ab"
    archived = tmp_path / ".trash" / "0123456789ab-old"
    active.mkdir(parents=True)
    archived.mkdir(parents=True)
    document = {
        "id": "0123456789ab",
        "segments": [{"translation": "partial translation"}],
    }
    (active / "document.json").write_text(json.dumps(document), encoding="utf-8")
    (archived / "document.json").write_text(json.dumps(document), encoding="utf-8")
    (active / "document.json.tmp").write_text("temporary", encoding="utf-8")
    exported = []

    monkeypatch.setattr(storage_module, "WORKSPACE", tmp_path)
    monkeypatch.setattr(storage_module, "create_translated_docx", lambda item: exported.append(item["id"]))

    removed, export_count = storage_module.finalize_and_clear_document_records()

    assert export_count == 1
    assert exported == ["0123456789ab"]
    assert removed == 3
    assert not list(tmp_path.rglob("document.json"))
    assert not list(tmp_path.rglob("document.json.tmp"))


class _FakeResponse:
    def __init__(self, status_code, data=None, text="", headers=None):
        self.status_code = status_code
        self._data = data or {}
        self.text = text
        self.headers = headers or {}

    def json(self):
        return self._data


def _translation_settings(max_retries=5):
    return SimpleNamespace(
        translation_configured=True,
        base_url="https://gateway.example/v1",
        protocol="openai",
        use_system_proxy=False,
        max_retries=max_retries,
        model="test-model",
        api_key="",
    )


def test_transient_model_failures_retry_then_continue(monkeypatch):
    success = {
        "choices": [
            {
                "message": {
                    "content": json.dumps(
                        {"translations": [{"id": "s1", "text": "translated"}]}
                    )
                }
            }
        ]
    }
    responses = [
        _FakeResponse(503, text="temporarily unavailable"),
        _FakeResponse(429, text="rate limited", headers={"Retry-After": "0"}),
        _FakeResponse(200, data=success),
    ]
    sleeps = []
    notices = []

    class FakeClient:
        def __init__(self, *args, **kwargs):
            pass

        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return False

        async def post(self, *args, **kwargs):
            return responses.pop(0)

    async def fake_sleep(seconds):
        sleeps.append(seconds)

    async def on_retry(notice):
        notices.append(notice)

    monkeypatch.setattr(translator_module.httpx, "AsyncClient", FakeClient)
    monkeypatch.setattr(translator_module.asyncio, "sleep", fake_sleep)
    monkeypatch.setattr(translator_module.random, "uniform", lambda *args: 0)

    result = asyncio.run(
        translate_batch(
            [{"id": "s1", "text": "source"}],
            "en",
            "zh",
            [],
            "",
            _translation_settings(),
            on_retry=on_retry,
        )
    )

    assert result == {"s1": "translated"}
    assert [notice["retry_number"] for notice in notices] == [1, 2]
    assert sleeps == [2.0, 4.0]
    assert not responses


def test_permanent_model_error_does_not_retry(monkeypatch):
    responses = [_FakeResponse(403, text="forbidden")]
    sleeps = []

    class FakeClient:
        def __init__(self, *args, **kwargs):
            pass

        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return False

        async def post(self, *args, **kwargs):
            return responses.pop(0)

    async def fake_sleep(seconds):
        sleeps.append(seconds)

    monkeypatch.setattr(translator_module.httpx, "AsyncClient", FakeClient)
    monkeypatch.setattr(translator_module.asyncio, "sleep", fake_sleep)

    with pytest.raises(TranslationError, match="HTTP 403"):
        asyncio.run(
            translate_batch(
                [{"id": "s1", "text": "source"}],
                "en",
                "zh",
                [],
                "",
                _translation_settings(),
            )
        )

    assert sleeps == []
    assert not responses

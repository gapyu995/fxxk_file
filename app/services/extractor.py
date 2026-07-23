from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".doc", ".docx", ".pdf"}


class ExtractionError(RuntimeError):
    pass


def extract_paragraphs(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ExtractionError(f"不支持 {suffix or '无扩展名'} 文件；请使用 DOC、DOCX 或 PDF。")
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    return _extract_legacy_doc(path)


def _extract_docx(path: Path) -> list[str]:
    try:
        document = Document(path)
    except Exception as exc:
        raise ExtractionError(f"DOCX 文件无法读取：{exc}") from exc

    # Only text-bearing paragraphs become translation units. Paragraphs that
    # contain images but no text are intentionally skipped and remain intact
    # when the translated DOCX is produced from the original template.
    blocks = [paragraph.text.strip() for paragraph in iter_docx_text_paragraphs(document)]
    if not blocks:
        raise ExtractionError("DOCX 中没有发现可提取的文字。")
    return blocks


def _extract_pdf(path: Path) -> list[str]:
    try:
        reader = PdfReader(str(path))
        blocks: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            blocks.extend(_paragraphs_from_text(text))
    except Exception as exc:
        raise ExtractionError(f"PDF 文件无法读取：{exc}") from exc
    if not blocks:
        raise ExtractionError("PDF 中没有可提取文字；它可能是扫描件，需要先进行 OCR。")
    return blocks


def _extract_legacy_doc(path: Path) -> list[str]:
    with tempfile.TemporaryDirectory(prefix="chanslator-") as temp_dir:
        converted = convert_legacy_doc(path, Path(temp_dir))
        return _extract_docx(converted)


def find_soffice() -> str | None:
    found = shutil.which("soffice") or shutil.which("libreoffice")
    if found:
        return found
    candidates = (
        Path("C:/Program Files/LibreOffice/program/soffice.exe"),
        Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
    )
    return str(next((item for item in candidates if item.exists()), "")) or None


def convert_legacy_doc(path: Path, output_dir: Path) -> Path:
    soffice = find_soffice()
    if not soffice:
        raise ExtractionError("读取旧版 .doc 需要安装 LibreOffice。也可以先在 Word 中另存为 .docx。")
    output_dir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "docx", "--outdir", str(output_dir), str(path)],
        capture_output=True,
        text=True,
        timeout=120,
        check=False,
    )
    converted = output_dir / f"{path.stem}.docx"
    if result.returncode != 0 or not converted.exists():
        detail = (result.stderr or result.stdout).strip()
        raise ExtractionError(f"LibreOffice 转换 .doc 失败：{detail or '未知错误'}")
    return converted


def iter_docx_text_paragraphs(document):
    """Yield all non-empty body/table paragraphs in a reproducible order."""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            yield paragraph
    seen_cells: set[int] = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        yield paragraph


def _paragraphs_from_text(text: str) -> list[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs: list[str] = []
    current: list[str] = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            if current:
                paragraphs.append(" ".join(current))
                current = []
        else:
            current.append(stripped)
    if current:
        paragraphs.append(" ".join(current))
    return paragraphs

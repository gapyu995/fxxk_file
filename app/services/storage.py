from __future__ import annotations

import json
import re
import shutil
import subprocess
import threading
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt

from app.config import ROOT
from app.services.extractor import convert_legacy_doc, find_soffice, iter_docx_text_paragraphs


INBOX = ROOT / "inbox"
ORIGINALS = ROOT / "originals"
WORKSPACE = ROOT / "workspace"
OUTPUT = ROOT / "output"
GLOSSARIES = ROOT / "glossaries"
STATIC = ROOT / "app" / "static"

_locks: dict[str, threading.RLock] = {}
_master_lock = threading.Lock()


def ensure_directories() -> None:
    for folder in (INBOX, ORIGINALS, WORKSPACE, OUTPUT, GLOSSARIES):
        folder.mkdir(parents=True, exist_ok=True)


def document_lock(document_id: str) -> threading.RLock:
    with _master_lock:
        return _locks.setdefault(document_id, threading.RLock())


def safe_filename(filename: str) -> str:
    name = Path(filename).name.strip().replace("\x00", "")
    name = re.sub(r"[<>:\"/\\|?*]", "_", name)
    return name[:180] or "document"


def document_dir(document_id: str) -> Path:
    return WORKSPACE / document_id


def save_document(document: dict[str, Any], render_files: bool = True) -> None:
    document["updated_at"] = utc_now()
    folder = document_dir(document["id"])
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / "document.json"
    temp = folder / "document.json.tmp"
    with document_lock(document["id"]):
        temp.write_text(json.dumps(document, ensure_ascii=False, indent=2), encoding="utf-8")
        temp.replace(path)


def load_document(document_id: str) -> dict[str, Any]:
    if not re.fullmatch(r"[a-f0-9]{12}", document_id):
        raise FileNotFoundError(document_id)
    path = document_dir(document_id) / "document.json"
    with document_lock(document_id):
        return json.loads(path.read_text(encoding="utf-8"))


def list_documents() -> list[dict[str, Any]]:
    documents: list[dict[str, Any]] = []
    if not WORKSPACE.exists():
        return documents
    for path in WORKSPACE.glob("*/document.json"):
        try:
            doc = json.loads(path.read_text(encoding="utf-8"))
            documents.append(
                {
                    "id": doc["id"],
                    "name": doc["name"],
                    "status": doc.get("status", "ready"),
                    "progress": doc.get("progress", 0),
                    "source_language": doc.get("source_language", "zh"),
                    "target_language": doc.get("target_language", "en"),
                    "updated_at": doc.get("updated_at", ""),
                }
            )
        except (OSError, json.JSONDecodeError, KeyError):
            continue
    return sorted(documents, key=lambda item: item["updated_at"], reverse=True)


def archive_document_record(document_id: str) -> Path:
    """Remove a task from recent documents without deleting original/output files."""
    if not re.fullmatch(r"[a-f0-9]{12}", document_id):
        raise FileNotFoundError(document_id)
    source = document_dir(document_id)
    if not source.is_dir():
        raise FileNotFoundError(document_id)
    trash = WORKSPACE / ".trash"
    trash.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = trash / f"{document_id}-{stamp}"
    number = 2
    while target.exists():
        target = trash / f"{document_id}-{stamp}-{number}"
        number += 1
    shutil.move(str(source), target)
    return target


def finalize_and_clear_document_records() -> tuple[int, int]:
    """Write the latest partial translations, then remove all document JSON state."""
    exported = 0
    if not WORKSPACE.exists():
        return 0, 0

    # A translation may be stopped while a large document is still in progress.
    # Preserve all text already returned by the model before removing its state.
    for path in WORKSPACE.glob("*/document.json"):
        try:
            document = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, ValueError, KeyError):
            continue
        if any(segment.get("translation", "").strip() for segment in document.get("segments", [])):
            try:
                create_translated_docx(document)
                exported += 1
            except Exception:
                # Closing must continue even when an old DOC or damaged record
                # cannot be exported. Existing originals and output files remain.
                pass

    removed = 0
    record_paths = set(WORKSPACE.rglob("document.json"))
    record_paths.update(WORKSPACE.rglob("document.json.tmp"))
    for path in record_paths:
        try:
            path.unlink(missing_ok=True)
            removed += 1
        except OSError:
            continue
    return removed, exported


def create_docx_exports(document: dict[str, Any]) -> tuple[Path, Path]:
    stem = safe_filename(Path(document["name"]).stem)
    bilingual_path = OUTPUT / f"{stem}_{document['id']}_双语对照.docx"
    translated_path = OUTPUT / f"{stem}_{document['id']}_译文.docx"

    bilingual = Document()
    bilingual.add_heading(f"{stem} — 双语对照稿", level=1)
    table = bilingual.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    headers = table.rows[0].cells
    headers[0].text = "原文"
    headers[1].text = "译文"
    for segment in document.get("segments", []):
        cells = table.add_row().cells
        cells[0].text = segment.get("source", "")
        cells[1].text = segment.get("translation", "")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
    bilingual.save(bilingual_path)

    create_translated_docx(document, translated_path)
    return bilingual_path, translated_path


def original_path(document: dict[str, Any]) -> Path:
    relative = Path(document["original_path"])
    path = (ROOT / relative).resolve()
    path.relative_to(ROOT.resolve())
    if not path.is_file():
        raise FileNotFoundError(path)
    return path


def translated_output_path(document: dict[str, Any]) -> Path:
    stem = safe_filename(Path(document["name"]).stem)
    return OUTPUT / f"{stem}_{document['id']}_译文.docx"


def create_translated_docx(document: dict[str, Any], target: Path | None = None) -> Path:
    """Create a standalone translation while retaining DOC/DOCX images/layout."""
    target = target or translated_output_path(document)
    target.parent.mkdir(parents=True, exist_ok=True)
    source = original_path(document)
    suffix = source.suffix.lower()
    segments = document.get("segments", [])

    if suffix in {".doc", ".docx"}:
        template = source
        if suffix == ".doc":
            template_dir = document_dir(document["id"]) / "source-template"
            expected = template_dir / f"{source.stem}.docx"
            template = expected if expected.exists() else convert_legacy_doc(source, template_dir)
        shutil.copy2(template, target)
        translated = Document(target)
        paragraphs = list(iter_docx_text_paragraphs(translated))
        if len(paragraphs) != len(segments):
            raise RuntimeError(
                f"无法按原版面回填译文：原文有 {len(paragraphs)} 个文本段，工作区有 {len(segments)} 个翻译段。"
            )
        for paragraph, segment in zip(paragraphs, segments):
            _replace_text_preserving_drawings(paragraph, segment.get("translation", "").strip())
    else:
        translated = Document()
        translated.core_properties.title = f"{document['name']} - Translation"
        translated.add_heading(Path(document["name"]).stem, level=1)
        for segment in segments:
            text = segment.get("translation", "").strip()
            if text:
                translated.add_paragraph(text)
    translated.save(target)
    return target


def create_original_preview(document: dict[str, Any]) -> Path | None:
    """Return the original PDF or a cached PDF rendering of DOC/DOCX."""
    source = original_path(document)
    if source.suffix.lower() == ".pdf":
        return source
    if source.suffix.lower() not in {".doc", ".docx"}:
        return None
    soffice = find_soffice()
    if not soffice:
        return None
    folder = document_dir(document["id"]) / "preview"
    folder.mkdir(parents=True, exist_ok=True)
    target = folder / "original.pdf"
    if target.exists() and target.stat().st_mtime >= source.stat().st_mtime:
        return target
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(folder), str(source)],
        capture_output=True,
        text=True,
        timeout=180,
        check=False,
    )
    generated = folder / f"{source.stem}.pdf"
    if result.returncode != 0 or not generated.exists():
        return None
    if generated != target:
        generated.replace(target)
    return target


def _replace_text_preserving_drawings(paragraph, value: str) -> None:
    """Replace w:t nodes only, leaving drawings/pictures and paragraph style intact."""
    text_nodes = paragraph._p.xpath(".//w:t")
    if text_nodes:
        text_nodes[0].text = value
        for node in text_nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(value)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")
